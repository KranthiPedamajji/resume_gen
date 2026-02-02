from pathlib import Path
from typing import Dict, List
import logging
import re
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

from app.models.schemas import ResumeState


_PLACEHOLDER_MAP = {
    "PROFESSIONAL SUMMARY": "{{PROFESSIONAL_SUMMARY}}",
    "TECHNICAL SKILLS": "{{TECHNICAL_SKILLS}}",
    "PROFESSIONAL EXPERIENCE": "{{PROFESSIONAL_EXPERIENCE}}",
    "EDUCATION": "{{EDUCATION}}",
}


_DEF_FONT = "Calibri"
_DEF_FONT_SIZE = 12
_BLACK = RGBColor(0, 0, 0)

_BULLET_PREFIX_RE = re.compile(r"^\s*(?:[-•*]|\d+\.)\s+")
_DATE_RANGE_RE = re.compile(
    r"\b(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|"
    r"Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|"
    r"Dec|December)\s+\d{4}\s*[–-]\s*(?:Present|"
    r"(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|"
    r"Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|"
    r"Dec|December)\s+\d{4})\b",
    re.IGNORECASE,
)
_RIGHT_TAB_STOP_INCH = 6.5
_MAX_DOCX_BYTES = 1_800_000

logger = logging.getLogger(__name__)


def sanitize_name(name: str) -> str:
    """Sanitize a string for filesystem-safe paths."""
    cleaned = re.sub(r"[\\/:*?\"<>|]", "", name)
    cleaned = cleaned.replace(" ", "_")
    cleaned = re.sub(r"_+", "_", cleaned)
    return cleaned.strip("_") or "UNKNOWN"


def build_output_paths(company_name: str, position_name: str, job_id: str | None) -> Path:
    """Build the output directory path for exported resumes."""
    company = sanitize_name(company_name)
    position = sanitize_name(position_name)
    suffix = sanitize_name(job_id) if job_id else "NOJOBID"
    folder_name = f"{position}_{suffix}"
    return Path(r"S:\applications") / company / folder_name


def parse_sections_from_resume_text(resume_text: str) -> Dict[str, List[str]]:
    """Parse resume text into sections based on headings."""
    lines = [line.strip() for line in resume_text.splitlines()]
    sections: Dict[str, List[str]] = {}
    current = None

    for line in lines:
        if not line:
            continue
        if _is_separator_line(line):
            continue
        heading = _normalize_heading(line)
        if heading:
            current = heading
            sections.setdefault(current, [])
            continue
        if current:
            sections[current].append(line)

    mapped = {}
    for key, values in sections.items():
        if key == "CORE SKILLS":
            mapped["TECHNICAL SKILLS"] = values
        elif key == "EXPERIENCE HIGHLIGHTS":
            mapped["PROFESSIONAL EXPERIENCE"] = values
        else:
            mapped[key] = values

    return mapped


def _normalize_heading(line: str) -> str | None:
    normalized = _clean_markdown(line.strip()).lstrip("#").strip()
    upper = normalized.upper()
    if upper in {
        "PROFESSIONAL SUMMARY",
        "CORE SKILLS",
        "TECHNICAL SKILLS",
        "EXPERIENCE HIGHLIGHTS",
        "PROFESSIONAL EXPERIENCE",
        "EDUCATION",
    }:
        return upper
    return None


def _clean_markdown(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^#+\s*", "", cleaned)
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("__", "")
    cleaned = cleaned.replace("*", "")
    return cleaned.strip()



def _relocate_education_block(doc: Document) -> None:
    """Ensure the EDUCATION block is at the end of the document."""
    paragraphs = list(doc.paragraphs)
    start_idx = None
    for idx, para in enumerate(paragraphs):
        para_text = para.text or ""
        if "{{EDUCATION}}" in para_text:
            start_idx = idx
            break
        if _clean_markdown(para_text).strip().upper() == "EDUCATION":
            start_idx = idx
            break
    if start_idx is None:
        return

    body = doc._element.body
    for para in paragraphs[start_idx:]:
        body.remove(para._p)
        body.append(para._p)



def _ensure_blank_before_education(doc: Document) -> None:
    """Insert a blank paragraph before the EDUCATION heading if needed."""
    paragraphs = list(doc.paragraphs)
    target = None
    for para in paragraphs:
        para_text = para.text or ""
        if "{{EDUCATION}}" in para_text:
            target = para
            break
        if _clean_markdown(para_text).strip().upper() == "EDUCATION":
            target = para
            break
    if target is None:
        return
    prev = target._p.getprevious()
    if prev is not None:
        prev_para = None
        for p in doc.paragraphs:
            if p._p is prev:
                prev_para = p
                break
        if prev_para is not None and (prev_para.text or '').strip() == '':
            return
    blank = doc.add_paragraph("")
    target._p.addprevious(blank._p)


def _strip_embedded_font_nodes(font_table_xml: bytes) -> bytes:
    """Remove embedded font tags while preserving original namespace prefixes."""
    try:
        xml = font_table_xml.decode("utf-8")
    except UnicodeDecodeError:
        return font_table_xml

    # Preserve prefix mappings/MC attributes; strip only embed tag elements.
    xml = re.sub(
        r"<w:embed(?:Regular|Bold|Italic|BoldItalic)\b[^>]*/>",
        "",
        xml,
        flags=re.IGNORECASE,
    )
    return xml.encode("utf-8")


def _strip_font_relationships(font_rels_xml: bytes) -> bytes:
    """Remove Relationship nodes that target embedded font binaries."""
    try:
        xml = font_rels_xml.decode("utf-8")
    except UnicodeDecodeError:
        return font_rels_xml

    xml = re.sub(
        r'<Relationship\b[^>]*\bTarget=["\']fonts/[^"\']+["\'][^>]*/>',
        "",
        xml,
        flags=re.IGNORECASE,
    )
    return xml.encode("utf-8")


def _optimize_docx_file(docx_path: Path, max_bytes: int = _MAX_DOCX_BYTES) -> None:
    """Shrink DOCX size by stripping embedded fonts when above threshold."""
    try:
        if not docx_path.exists() or docx_path.stat().st_size <= max_bytes:
            return
    except OSError:
        return

    with tempfile.NamedTemporaryFile(suffix=".docx", dir=str(docx_path.parent), delete=False) as tmp_file:
        tmp_path = Path(tmp_file.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
            tmp_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as dst:
            for info in src.infolist():
                name = info.filename
                if name.startswith("word/fonts/"):
                    continue

                data = src.read(name)
                if name == "word/fontTable.xml":
                    data = _strip_embedded_font_nodes(data)
                elif name == "word/_rels/fontTable.xml.rels":
                    data = _strip_font_relationships(data)

                dst.writestr(name, data)

        original_size = docx_path.stat().st_size
        optimized_size = tmp_path.stat().st_size if tmp_path.exists() else original_size
        if optimized_size < original_size:
            tmp_path.replace(docx_path)
            logger.info(
                "Optimized DOCX %s: %d -> %d bytes",
                docx_path.name,
                original_size,
                optimized_size,
            )
        else:
            tmp_path.unlink(missing_ok=True)
    except Exception as exc:
        logger.warning("DOCX optimization skipped for %s: %s", docx_path, exc)
        tmp_path.unlink(missing_ok=True)



def export_resume_to_docx(template_path: Path, sections: Dict[str, List[str]], output_path: Path) -> None:
    """Render resume sections into the DOCX template and save it."""
    doc = Document(template_path)
    _relocate_education_block(doc)
    _ensure_blank_before_education(doc)

    for paragraph in list(doc.paragraphs):
        for key, placeholder in _PLACEHOLDER_MAP.items():
            if placeholder in paragraph.text:
                if key == "EDUCATION":
                    continue
                content = sections.get(key, [])
                _replace_placeholder_in_paragraph(doc, paragraph, placeholder, content, key)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    _optimize_docx_file(output_path)


def export_docx_from_state(state: ResumeState, template_path: Path, output_path: Path) -> None:
    """Render a ResumeState into the DOCX template and save it."""
    sections = _sections_from_state(state)
    export_resume_to_docx(template_path, sections, output_path)


def _sections_from_state(state: ResumeState) -> Dict[str, List[str]]:
    summary_lines = state.sections.professional_summary.splitlines() if state.sections.professional_summary else []
    skills_lines = state.sections.technical_skills or []

    experience_lines: List[str] = []
    for role in state.sections.experience:
        header = _format_role_header(role)
        experience_lines.append(header)
        for bullet in role.bullets:
            experience_lines.append(bullet)

    return {
        "PROFESSIONAL SUMMARY": summary_lines,
        "TECHNICAL SKILLS": skills_lines,
        "PROFESSIONAL EXPERIENCE": experience_lines,
    }


def _format_role_header(role) -> str:
    parts = [role.company]
    if role.title:
        parts.append(role.title)
    header = " - ".join(parts) if len(parts) > 1 else parts[0]
    tail = []
    if role.location:
        tail.append(role.location)
    if role.dates:
        tail.append(role.dates)
    if tail:
        header = f"{header} | " + " | ".join(tail)
    return header


def _replace_placeholder_in_paragraph(
    doc: Document,
    paragraph,
    placeholder: str,
    content: List[str],
    section_key: str,
) -> None:
    """Replace a placeholder in a paragraph and insert additional lines preserving style."""
    if not content:
        if section_key == "EDUCATION":
            paragraph.text = paragraph.text.replace(placeholder, "Education available upon request")
        else:
            paragraph.text = paragraph.text.replace(placeholder, "")
        _apply_run_format(paragraph)
        return

    first_line = content[0]
    paragraph.text = paragraph.text.replace(placeholder, "")
    paragraph.style = _resolve_style(doc, paragraph.style, first_line, section_key)
    base_style = paragraph.style

    if section_key == "EDUCATION":
        _apply_run_format(paragraph)
        insert_after = _insert_paragraph_after(
            doc,
            paragraph,
            first_line,
            base_style,
            section_key,
        )
        remaining = content[1:]
    else:
        insert_after = _set_paragraph_content(paragraph, first_line, section_key, doc)
        _apply_run_format(paragraph)
        remaining = content[1:]

    for line in remaining:
        if section_key == "PROFESSIONAL EXPERIENCE" and _is_role_header(_clean_markdown(line)):
            insert_after = _insert_blank_paragraph_after(doc, insert_after)
        insert_after = _insert_paragraph_after(
            doc,
            insert_after,
            line,
            base_style,
            section_key,
        )


def _insert_paragraph_after(doc: Document, paragraph, text: str, style, section_key: str) -> any:
    new_para = doc.add_paragraph(text)
    new_para.style = _resolve_style(doc, style, text, section_key)
    last_para = _set_paragraph_content(new_para, text, section_key, doc)
    _apply_run_format(new_para)
    paragraph._p.addnext(new_para._p)
    return last_para


def _insert_blank_paragraph_after(doc: Document, paragraph) -> any:
    new_para = doc.add_paragraph("")
    paragraph._p.addnext(new_para._p)
    return new_para


def _normalize_line_for_style(text: str, style) -> str:
    """Normalize bullet markers when inserting into list-style paragraphs."""
    if not style or not getattr(style, "name", ""):
        return text
    style_name = style.name.lower()
    if "list" in style_name or "bullet" in style_name:
        return text.lstrip("- ").lstrip("• ").strip()
    return text


def _resolve_style(doc: Document, base_style, text: str, section_key: str):
    if _should_be_bullet(text, section_key):
        return _get_bullet_style(doc, base_style)
    return base_style


def _get_bullet_style(doc: Document, base_style):
    if base_style and getattr(base_style, "name", ""):
        style_name = base_style.name.lower()
        if "list" in style_name or "bullet" in style_name:
            return base_style
    for candidate in ["List Bullet", "List Bullet 2", "List Paragraph"]:
        if candidate in doc.styles:
            return doc.styles[candidate]
    return base_style


def _set_paragraph_content(paragraph, text: str, section_key: str, doc: Document) -> None:
    paragraph.text = ""
    cleaned = _clean_markdown(text)
    is_bullet = _should_be_bullet(cleaned, section_key)
    if _is_bullet_line(cleaned):
        cleaned = _strip_bullet_prefix(cleaned)
    if is_bullet and not _is_bullet_style(paragraph.style):
        cleaned = f"• {cleaned}"

    if section_key == "TECHNICAL SKILLS" and ":" in cleaned:
        label, rest = cleaned.split(":", 1)
        _add_run(paragraph, f"{label.strip()}: ", bold=True)
        _add_run(paragraph, rest.strip(), bold=False)
        _apply_run_format(paragraph)
        return paragraph

    if section_key == "PROFESSIONAL EXPERIENCE" and not is_bullet and _is_role_header(cleaned):
        company, title, location, dates = _parse_role_header(cleaned)
        _set_company_date_line(paragraph, company, dates)
        _apply_run_format(paragraph)
        detail = " | ".join([part for part in [title, location] if part])
        if detail:
            return _insert_plain_paragraph_after(doc, paragraph, detail, paragraph.style, bold=True, color=_BLACK)
        return paragraph

    _add_run(paragraph, cleaned, bold=False)
    _apply_run_format(paragraph)
    return paragraph


def _add_run(paragraph, text: str, bold: bool, color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _is_bullet_line(text: str) -> bool:
    return bool(_BULLET_PREFIX_RE.match(text))


def _strip_bullet_prefix(text: str) -> str:
    return _BULLET_PREFIX_RE.sub("", text).strip()


def _is_role_header(text: str) -> bool:
    if _DATE_RANGE_RE.search(text):
        return True
    if re.search(r"\b(19|20)\d{2}\b", text) and (" | " in text or "—" in text or " - " in text):
        return True
    return False


def _is_separator_line(text: str) -> bool:
    cleaned = _clean_markdown(text).strip()
    if not cleaned:
        return True
    return bool(re.fullmatch(r"[-–—]{2,}", cleaned))


def _parse_role_header(text: str) -> tuple[str, str | None, str | None, str | None]:
    match = _DATE_RANGE_RE.search(text)
    dates = None
    base = text
    if match:
        dates = match.group(0).replace("\u2013", "-").replace("\u2014", "-")
        base = (text[:match.start()] + text[match.end():]).strip(" -|,")

    company = base
    title = None
    location = None

    if " - " in base:
        company, rest = base.split(" - ", 1)
        parts = [p.strip() for p in rest.split("|") if p.strip()]
        if parts:
            title = parts[0]
        if len(parts) > 1:
            location = parts[1]
    else:
        parts = [p.strip() for p in base.split("|") if p.strip()]
        if parts:
            company = parts[0]
        if len(parts) > 1:
            title = parts[1]
        if len(parts) > 2:
            location = parts[2]

    return company.strip(), title, location, dates


def _set_company_date_line(paragraph, company: str, dates: str | None) -> None:
    paragraph.text = ""
    if dates:
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(_RIGHT_TAB_STOP_INCH),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.SPACES,
        )
        run = paragraph.add_run(company.strip())
        run.bold = True
        paragraph.add_run("\t")
        run_dates = paragraph.add_run(dates.strip())
        run_dates.bold = True
    else:
        run = paragraph.add_run(company.strip())
        run.bold = True


def _insert_plain_paragraph_after(doc: Document, paragraph, text: str, style, bold: bool = False, color: RGBColor | None = None):
    new_para = doc.add_paragraph("")
    new_para.style = style
    _add_run(new_para, text, bold=bold, color=color)
    _apply_run_format(new_para)
    paragraph._p.addnext(new_para._p)
    return new_para


def _is_bullet_style(style) -> bool:
    if not style or not getattr(style, "name", ""):
        return False
    style_name = style.name.lower()
    return "list" in style_name or "bullet" in style_name


def _should_be_bullet(text: str, section_key: str) -> bool:
    if section_key == "TECHNICAL SKILLS":
        return True
    if section_key == "PROFESSIONAL EXPERIENCE":
        cleaned = _clean_markdown(text)
        return not _is_role_header(cleaned)
    return _is_bullet_line(text)


def _apply_run_format(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = _DEF_FONT
        run.font.size = Pt(_DEF_FONT_SIZE)
